
from flask import Flask, render_template, request, redirect, url_for, session
from flask_sqlalchemy import SQLAlchemy
from utils.image_processor import process_and_save_image
from sqlalchemy import func
from werkzeug.utils import secure_filename
from datetime import datetime
import os
import uuid
#from flask_wtf import CSRFProtect
import os
import csv
import io
import zipfile
from flask import Response
from werkzeug.security import generate_password_hash, check_password_hash
#from flask import flash
import shutil
import os
from datetime import datetime
from flask import send_file
from itsdangerous import URLSafeTimedSerializer
#from your_file_name import InternProfile
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from uuid import uuid4
from utils.image_processor import (
    process_and_save_image,
    process_and_overwrite_cropped_image
)
from flask import request, session
import cloudinary
import cloudinary.uploader
import os
import cloudinary
import cloudinary.uploader
from io import BytesIO
from PIL import Image
import base64
cloudinary.config(
    cloud_name=os.environ.get("CLOUDINARY_CLOUD_NAME"),
    api_key=os.environ.get("CLOUDINARY_API_KEY"),
    api_secret=os.environ.get("CLOUDINARY_API_SECRET"),
    secure=True
)
import pillow_heif
pillow_heif.register_heif_opener()

pdfmetrics.registerFont(
    TTFont('CinzelBlack', 'fonts/CinzelDecorative-Black.ttf')
)

pdfmetrics.registerFont(
    TTFont('PatrickRegular', 'fonts/PatrickHand-Regular.ttf')
)

from flask import send_file

from flask_bcrypt import Bcrypt

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.styles import getSampleStyleSheet
from datetime import date
import os
from flask_migrate import Migrate

BASE_DIR = os.path.abspath(os.path.dirname(__file__))

from flask_login import login_required, current_user

from flask import Flask, render_template, request, redirect, url_for, session, flash

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024  # 50 MB
print("MAX LIMIT:", app.config.get("MAX_CONTENT_LENGTH"))

app.config["SECRET_KEY"] = "bbp-super-secure-key-2026"
bcrypt = Bcrypt(app)
# 🔥 TEMPORARY: Disable CSRF for development
app.config['WTF_CSRF_ENABLED'] = False
serializer = URLSafeTimedSerializer(app.config["SECRET_KEY"])
app.config["CERTIFICATE_FOLDER"] = os.path.join(
    app.root_path,
    "storage",
    "certificates"
)
os.makedirs(app.config["CERTIFICATE_FOLDER"], exist_ok=True)
app.config["REPORT_FOLDER"] = os.path.join(BASE_DIR, "uploads", "reports")
os.makedirs(app.config["REPORT_FOLDER"], exist_ok=True)


BACKUP_FOLDER = os.path.join(os.getcwd(), "backups")

if not os.path.exists(BACKUP_FOLDER):
    os.makedirs(BACKUP_FOLDER)

# =========================
# BASIC CONFIG
# =========================

#csrf = CSRFProtect(app)

@app.route("/test-public")
def test_public():
    return "Public Works"

import os

DATABASE_URL = os.environ.get("DATABASE_URL")

if DATABASE_URL:
    # Render gives postgres:// which SQLAlchemy dislikes
    if DATABASE_URL.startswith("postgres://"):
        DATABASE_URL = DATABASE_URL.replace(
            "postgres://",
            "postgresql://",
            1
        )

    app.config["SQLALCHEMY_DATABASE_URI"] = DATABASE_URL
else:
    # Local development fallback
    app.config["SQLALCHEMY_DATABASE_URI"] = "postgresql://localhost/bbp_db"

app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

print("📂 Using database:", app.config["SQLALCHEMY_DATABASE_URI"])


# =========================
# DATABASE INIT
# =========================
db = SQLAlchemy(app)
migrate = Migrate(app, db)
CONSERVATION_STATUSES = [
    "Critically Endangered",
    "Endangered",
    "Vulnerable",
    "Near Threatened",
    "Least Concern"
]

HABITATS = [
    "Forest",
    "Grassland",
    "Wetlands",
    "Scrubland",
    "Rocky Terrain",
    "Mixed Habitat"
]
TARGET_AUDIENCES = [
    "School Students",
    "College Students",
    "Teachers",
    "Families",
    "General Public",
    "Researchers"
]






# =========================
# DATABASE MODELS
# =========================
class Event(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(100), nullable=False)
    description = db.Column(db.Text, nullable=False)
    event_date = db.Column(db.Date, nullable=False)
    cover_image = db.Column(db.String(500))   # 👈 ADD THIS
    image = db.Column(db.String(500))
    icon = db.Column(db.String(255), nullable=True)

class EventImage(db.Model):
    __tablename__ = "event_image"

    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(500), nullable=False)

    event_id = db.Column(
        db.Integer,
        db.ForeignKey("event.id"),
        nullable=False
    )

    event = db.relationship(
        "Event",
        backref=db.backref("images", lazy=True)
    )

    

class Program(db.Model):
    id = db.Column(db.Integer, primary_key=True)

    name = db.Column(db.String(150), nullable=False)  # Upcoming Awareness Day
    description = db.Column(db.Text, nullable=False)

    target_audience = db.Column(db.String(200), nullable=False)

    location = db.Column(db.String(200))

    event_date = db.Column(db.Date)

    start_time = db.Column(db.Time)
    end_time = db.Column(db.Time)

    category = db.Column(db.String(50))
    image = db.Column(db.String(200))

class Gallery(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    image = db.Column(db.String(255), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class Animal(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    common_name = db.Column(db.String(100), nullable=False)
    scientific_name = db.Column(db.String(150), nullable=False)
    habitat = db.Column(db.String(150), nullable=False)
    conservation_status = db.Column(db.String(50), nullable=False)
    description = db.Column(db.Text, nullable=False)
    image = db.Column(db.String(500))
class Admin(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)

    role = db.Column(db.String(20), default="admin")  
    # values: superadmin, admin

    is_approved = db.Column(db.Boolean, default=False)

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100))
    email = db.Column(db.String(120), unique=True)
    password_hash = db.Column(db.String(200))
    role = db.Column(db.String(20))  # volunteer / intern
    status = db.Column(db.String(20), default="pending")
    #is_approved = db.Column(db.Boolean, default=False)
    intern_profile = db.relationship(
    "InternProfile",
    backref="user",
    uselist=False,
    cascade="all, delete-orphan"
)
    daily_logs = db.relationship(
    "InternDailyLog",
    backref="intern",
    cascade="all, delete-orphan"
)

    reset_token = db.Column(db.String(200), nullable=True)

class InternProfile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    course = db.Column(db.String(150))
    year = db.Column(db.String(50))
    semester = db.Column(db.String(50))
    college = db.Column(db.String(200))
    phone = db.Column(db.String(20))
    start_date = db.Column(db.Date)
    end_date = db.Column(db.Date)
    objective = db.Column(db.Text)
    application_status = db.Column(db.String(50), default="submitted")
    report_upload_enabled = db.Column(db.Boolean, default=False)
    report_file = db.Column(db.String(255))
    admin_note = db.Column(db.Text)
    is_completed = db.Column(db.Boolean, default=False)
    certificate_file = db.Column(db.String(200))
    certificate_requested = db.Column(db.Boolean, default=False)
    internship_status = db.Column(db.String(20), default="pending")
    completion_date = db.Column(db.Date)
    survey_rating = db.Column(db.Integer)
    survey_completed = db.Column(db.Boolean, default=False)
    survey_feedback = db.Column(db.Text)   # store survey answers (JSON/text)
    daily_log_enabled = db.Column(db.Boolean, default=True)
    internship_subjects = db.Column(db.String(300))
    certificate_id = db.Column(db.String(30), unique=True, nullable=True)
    issued_at = db.Column(db.DateTime, nullable=True, default=datetime.utcnow)
    certificate_serial = db.Column(db.Integer)
    report_submission_date = db.Column(db.Date)
    photo = db.Column(db.String(255))
    
    
class InternDailyLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    intern_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    date = db.Column(db.Date, nullable=False)
    login_time = db.Column(db.DateTime)
    logout_time = db.Column(db.DateTime)
    summary = db.Column(db.Text)
    location = db.Column(db.String(150))
    schedule = db.Column(db.String(200))
    submitted = db.Column(db.Boolean, default=False)
    previous_summary = db.Column(db.Text, nullable=True)
    previous_location = db.Column(db.String(150), nullable=True)
    previous_schedule = db.Column(db.String(200), nullable=True)
    def __repr__(self):
        return f"<User {self.email} ({self.role})>"

class Outreach(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text, nullable=False)
    image = db.Column(db.String(300), nullable=True)
    created_at = db.Column(db.DateTime, default=db.func.current_timestamp())
    views = db.Column(db.Integer, default=0)   
    
from datetime import datetime

class AuditLog(db.Model):
    __tablename__ = "audit_log"

    id = db.Column(db.Integer, primary_key=True)

    # Who performed action
    actor_id = db.Column(db.Integer, nullable=True)
    actor_role = db.Column(db.String(50), nullable=True)

    # What section (outreach, gallery, animals, interns, etc.)
    section = db.Column(db.String(100), nullable=False, index=True)

    # What action (create, update, delete, approve, etc.)
    action = db.Column(db.String(100), nullable=False, index=True)


    # What object was affected
    target_type = db.Column(db.String(100), nullable=True)
    target_id = db.Column(db.Integer, nullable=True)

    # Extra information
    description = db.Column(db.Text, nullable=True)

    # IP address
    ip_address = db.Column(db.String(100), nullable=True)

    # Timestamp
    timestamp = db.Column(db.DateTime, default=datetime.utcnow, nullable=False, index=True)

    def __repr__(self):
        return f"<AuditLog {self.section} - {self.action}>"



# =========================
# UPLOAD CONFIG
# =========================
ALLOWED_EXTENSIONS = {"jpg", "jpeg", "png", "pdf", "heic"}

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def save_image(file, folder):
    ext = file.filename.rsplit(".", 1)[1].lower()
    unique_name = f"{uuid.uuid4().hex}.{ext}"
    path = os.path.join(folder, unique_name)
    file.save(path)
    return unique_name


from functools import wraps
from flask_login import current_user
from flask import redirect, url_for, flash

def intern_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated:
            return redirect(url_for("login"))

        if current_user.role != "intern":
            flash("Access denied.", "danger")
            return redirect(url_for("home"))

        return f(*args, **kwargs)

    return decorated_function


from flask import session, redirect, url_for, flash

from flask import abort


from flask import session, redirect, url_for, flash


from functools import wraps
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get("admin_logged_in"):
            flash("Please login as admin.", "danger")
            return redirect(url_for("admin_login"))
        return f(*args, **kwargs)
    return decorated_function





def superadmin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if session.get("admin_role") != "superadmin":
            flash("Access denied.", "danger")
            return redirect(url_for("admin_dashboard"))
        return f(*args, **kwargs)
    return decorated_function

@app.route("/open")
def open_test():
    return "OPEN ACCESS"

# =========================
# PUBLIC ROUTES
# =========================
@app.route("/")
def home():
    print("HOME LOADED")
    print(session)
    return render_template("index.html")


@app.route("/test_limit", methods=["POST"])
def test_limit():
    print("MAX:", app.config['MAX_CONTENT_LENGTH'])
    return "OK"
from werkzeug.exceptions import RequestEntityTooLarge



@app.route("/contact")
def contact():
    return render_template("contact.html")

import os

@app.route("/gallery")
def public_gallery():

    images = Gallery.query.order_by(Gallery.id.desc()).all()

    return render_template("gallery.html", images=images)







from datetime import date



@app.route("/education-spotlights")
def education_spotlights():

    events = Event.query.order_by(Event.event_date.desc()).all()

    return render_template(
        "education_spotlights.html",
        events=events
    )


@app.route("/education-spotlight/<int:event_id>")
def spotlight_detail(event_id):

    event = Event.query.get_or_404(event_id)

    return render_template(
        "spotlight_detail.html",
        event=event
    )







from datetime import datetime

@app.route("/education-programs")
def education_programs():

    now = datetime.now()
    today = now.date()
    current_time = now.time()

    programs = Program.query.filter(
        (Program.event_date > today) |
        (
            (Program.event_date == today) &
            (Program.end_time >= current_time)
        )
    ).order_by(Program.event_date.asc()).all()

    return render_template(
        "programs_public.html",
        programs=programs
    )





from sqlalchemy import or_

from sqlalchemy import or_

@app.route("/animals")
def animals():
    search = request.args.get("search")
    status = request.args.get("status")

    query = Animal.query

    if search:
        query = query.filter(Animal.common_name.ilike(f"%{search}%"))

    if status and status != "all":
        query = query.filter(Animal.conservation_status == status)

    animals = query.all()

    return render_template(
        "animals.html",
        animals=animals,
        statuses=CONSERVATION_STATUSES
    )




# =========================
# ADMIN AUTH
# =========================
@app.route("/volunteer/login", methods=["GET", "POST"])
def volunteer_login():

    if request.method == "POST":

        email = request.form.get("email")
        password = request.form.get("password")

        user = User.query.filter_by(email=email, role="volunteer").first()

        if not user:
            flash("Account not found", "danger")
            return redirect(url_for("volunteer_login"))

        if not check_password_hash(user.password_hash, password):
            flash("Invalid password", "danger")
            return redirect(url_for("volunteer_login"))

        if user.status == "pending":
            flash("Your application is awaiting approval.", "warning")
            return redirect(url_for("volunteer_login"))
        if user.status == "rejected":
            flash("Your application was rejected.", "danger")
            return redirect(url_for("volunteer_login"))


        # ✅ SET SESSION WITHOUT CLEARING
        session["volunteer_id"] = user.id
        session["volunteer_name"] = user.name
        session.permanent = True

        return redirect(url_for("volunteer_dashboard"))

    return render_template("volunteer_login.html")

@app.route("/admin/login", methods=["GET", "POST"])
def admin_login():

    if session.get("admin_logged_in"):
        return redirect(url_for("admin_dashboard"))

    if request.method == "POST":

        username = request.form.get("username")
        password = request.form.get("password")

        admin = Admin.query.filter_by(username=username).first()

        # ❌ Admin not found
        if not admin:

            log_action(
                section="auth",
                action="failed_login",
                target_type="admin",
                target_id=None,
                description=f"Admin not found: {username}"
            )

            flash("Admin not found.", "danger")
            return render_template("admin_login.html", auth_page=True)

        # ❌ Wrong password
        if not bcrypt.check_password_hash(admin.password_hash, password):

            log_action(
                section="auth",
                action="failed_login",
                target_type="admin",
                target_id=admin.id,
                description=f"Invalid password attempt for: {username}"
            )

            flash("Invalid password.", "danger")
            return render_template("admin_login.html", auth_page=True)

        # ✅ SUCCESS LOGIN
        session["admin_logged_in"] = True
        session["admin_id"] = admin.id
        session["admin_username"] = admin.username
        session["admin_role"] = admin.role

        log_action(
            section="auth",
            action="login",
            target_type="admin",
            target_id=admin.id,
            description=f"Admin logged in: {username}"
        )

        flash("Login successful!", "success")
        return redirect(url_for("admin_dashboard"))

    return render_template("admin_login.html", auth_page=True)




@app.route("/admin/audit/cleanup")
@admin_required
def cleanup_logs():

    from datetime import timedelta
    cutoff = datetime.utcnow() - timedelta(days=180)

    AuditLog.query.filter(AuditLog.timestamp < cutoff).delete()
    db.session.commit()

    flash("Old logs deleted (older than 180 days).", "success")

    return redirect(url_for("admin_audit"))



@app.route("/admin/register", methods=["GET", "POST"])
def admin_register():

    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        confirm_password = request.form["confirm_password"]

        # Basic validations
        if password != confirm_password:
            flash("Passwords do not match", "danger")
            return redirect(url_for("admin_register"))

        # Check if username already exists
        existing_admin = Admin.query.filter_by(username=username).first()
        if existing_admin:
            flash("Username already exists", "danger")
            return redirect(url_for("admin_register"))

        # Hash password
        password_hash = generate_password_hash(password, method="pbkdf2:sha256")

        # Create admin
        new_admin = Admin(
            username=username,
            password_hash=password_hash
        )

        db.session.add(new_admin)
        db.session.commit()

        flash("Registration successful. Please login.", "success")
        return redirect(url_for("admin_login"))

    return render_template("admin_register.html")




# =========================
# ADMIN DASHBOARD
# =========================
@app.route("/admin/dashboard")
def admin_dashboard():

    total_volunteers = User.query.filter_by(
        role="volunteer",
        status="approved"
    ).count()

    total_interns = User.query.filter_by(
        role="intern",
        status="approved"
    ).count()

    pending_requests = User.query.filter_by(
        status="pending"
    ).count()

    certificate_requests = InternProfile.query.filter_by(
        certificate_requested=True,
        is_completed=False
    ).count()

    # 🔥 NEW COUNTS
    active_interns = InternProfile.query.filter_by(internship_status="active").count()
    completed_interns = InternProfile.query.filter_by(internship_status="completed").count()
    pending_internships = InternProfile.query.filter_by(internship_status="pending").count()

    return render_template(
        "admin_dashboard.html",
        total_volunteers=total_volunteers,
        total_interns=total_interns,
        pending_requests=pending_requests,
        certificate_requests=certificate_requests,
        active_interns=active_interns,
        completed_interns=completed_interns,
        pending_internships=pending_internships
    )







@app.route("/admin/volunteers")
def admin_volunteers():
    search = request.args.get("search")

    if search:
        volunteers = User.query.filter(
            User.role == "volunteer",
            (User.name.ilike(f"%{search}%") |
             User.email.ilike(f"%{search}%"))
        ).all()
    else:
        volunteers = User.query.filter_by(role="volunteer").all()

    return render_template(
        "admin_volunteers.html",
        volunteers=volunteers,
        search=search
    )

    
    
@app.route("/admin/interns")
def admin_interns():

    search = request.args.get("search")

    query = User.query.filter_by(role="intern", status="approved")

    if search:
        query = query.filter(User.name.ilike(f"%{search}%"))

    interns = query.all()

    return render_template("admin_interns.html", interns=interns)

@app.route("/admin/intern/<int:user_id>")
def admin_view_intern(user_id):

    intern = User.query.get_or_404(user_id)

    logs = InternDailyLog.query.filter_by(
        intern_id=user_id
    ).order_by(InternDailyLog.date).all()

    return render_template(
        "admin_intern_detail.html",
        intern=intern,
        logs=logs
    )




    
    
@app.route("/admin/users/delete/<int:user_id>")
@superadmin_required
def delete_user(user_id):
    user = User.query.get_or_404(user_id)

    if user.role == "superadmin":
        flash("Cannot remove Education Officer.", "danger")
        return redirect(url_for("admin_dashboard"))

    db.session.delete(user)
    db.session.commit()

    flash("User removed successfully.", "success")
    return redirect(request.referrer)






# =========================
# ADMIN EVENTS
# =========================
# =========================
# ADMIN EVENTS (CLOUDINARY VERSION)
# =========================
@app.route("/admin/events", methods=["GET", "POST"])
@admin_required
def admin_events():

    if request.method == "POST":

        # ================= BASIC EVENT DATA =================
        new_event = Event(
            title=request.form["title"],
            description=request.form["description"],
            event_date=datetime.strptime(
                request.form["event_date"],
                "%Y-%m-%d"
            )
        )

        db.session.add(new_event)
        db.session.flush()  # get event.id without full commit

        # ================= ICON UPLOAD =================
        icon_file = request.files.get("icon")

        if icon_file and icon_file.filename != "":
            upload_result = cloudinary.uploader.upload(
                icon_file,
                folder="bbp/events/icons",
                public_id=f"icon_{uuid4().hex}"
            )
            new_event.icon = upload_result["secure_url"]

        # ================= GALLERY IMAGES =================
        files = request.files.getlist("gallery_images")

        first_image_saved = False

        for file in files:
            if file and file.filename != "":
                upload_result = cloudinary.uploader.upload(
                    file,
                    folder="bbp/events/gallery",
                    public_id=f"event_{uuid4().hex}"
                )

                image_url = upload_result["secure_url"]

                event_image = EventImage(
                    filename=image_url,  # now storing full Cloudinary URL
                    event_id=new_event.id
                )

                db.session.add(event_image)

                # First image becomes cover
                if not first_image_saved:
                    new_event.cover_image = image_url
                    first_image_saved = True

        db.session.commit()

        log_action(
            section="event",
            action="create",
            target_type="event",
            target_id=new_event.id,
            description=f"Created event: {new_event.title}"
        )

        flash("Event added successfully.", "success")
        return redirect(url_for("admin_events"))

    events = Event.query.order_by(Event.event_date.desc()).all()

    return render_template(
        "admin_events.html",
        events=events
    )

@app.route("/admin/events/edit/<int:event_id>", methods=["GET", "POST"])
@admin_required
def edit_event(event_id):

    event = Event.query.get_or_404(event_id)

    if request.method == "POST":

        event.title = request.form["title"]
        event.description = request.form["description"]
        event.event_date = datetime.strptime(
            request.form["event_date"], "%Y-%m-%d"
        )

        # ================= ICON UPDATE (CLOUDINARY) =================
        icon_file = request.files.get("icon")

        if icon_file and icon_file.filename != "":
            upload_result = cloudinary.uploader.upload(
                icon_file,
                folder="bbp/events/icons",
                public_id=f"icon_{uuid4().hex}"
            )
            event.icon = upload_result["secure_url"]

        # ================= NEW GALLERY IMAGES (CLOUDINARY) =================
        images = request.files.getlist("gallery_images")

        for file in images:
            if file and file.filename != "":

                upload_result = cloudinary.uploader.upload(
                    file,
                    folder="bbp/events/gallery",
                    public_id=f"event_{uuid4().hex}"
                )

                image_url = upload_result["secure_url"]

                event_image = EventImage(
                    filename=image_url,
                    event_id=event.id
                )

                db.session.add(event_image)

                # If no cover exists, set first as cover
                if not event.cover_image:
                    event.cover_image = image_url

        db.session.commit()

        log_action(
            section="event",
            action="update",
            target_type="event",
            target_id=event.id,
            description=f"Updated event: {event.title}"
        )

        flash("Event updated successfully.", "success")
        return redirect(url_for("admin_events"))

    return render_template("edit_event.html", event=event)





@app.route("/admin/events/delete/<int:event_id>", methods=["POST"])
@admin_required
def delete_event(event_id):


    event = Event.query.get_or_404(event_id)

    event_title = event.title

    # Delete associated images
    for img in event.images:
        db.session.delete(img)

    db.session.delete(event)
    db.session.commit()

    log_action(
    section="event",
    action="delete",
    target_type="event",
    target_id=event_id,
    description=f"Deleted event: {event_title}"
)

    return redirect(url_for("admin_events"))

@app.route("/admin/crop-cover", methods=["POST"])
@admin_required
def crop_cover():

    file = request.files.get("cropped_image")

    if file:
        filename = f"{uuid.uuid4().hex}.jpg"
        save_path = os.path.join(app.static_folder, "images", "events", filename)
        file.save(save_path)

    return "", 200

@app.route("/admin/remove-user/<int:user_id>")
@admin_required
def remove_user(user_id):

    user = User.query.get_or_404(user_id)

    db.session.delete(user)
    db.session.commit()

    flash("User removed successfully.", "success")

    return redirect(request.referrer or url_for("admin_dashboard"))



from flask import request, redirect, url_for, render_template, flash


from werkzeug.utils import secure_filename
from PIL import Image as PILImage
import io
import base64
import time

# =========================
# ADMIN PROGRAMS
# =========================

@app.route("/admin/programs", methods=["GET", "POST"])
@admin_required
def admin_programs():

    today = date.today()
    view = request.args.get("view")

    upload_folder = os.path.join(
        app.root_path,
        "static",
        "uploads",
        "programs"
    )
    os.makedirs(upload_folder, exist_ok=True)

    if request.method == "POST":

        image_file = request.files.get("image")
        cropped_data = request.form.get("cropped_image")

        filename = None

        # ================= PRIORITY 1: CROPPED IMAGE =================
        if cropped_data and cropped_data.strip() != "":
            filename = f"{uuid4().hex}.webp"
            save_path = os.path.join(upload_folder, filename)

            process_and_overwrite_cropped_image(
                cropped_data,
                save_path
            )

        # ================= PRIORITY 2: NORMAL IMAGE =================
        elif image_file and image_file.filename != "":
            filename = process_and_save_image(
                image_file,
                upload_folder
            )

        # ================= TARGET AUDIENCE =================
        selected_audience = request.form.getlist("target_audience")
        other = request.form.get("other_audience")

        if other and other.strip():
            selected_audience.append(other.strip())

        if not selected_audience:
            flash("Please select at least one target audience.", "danger")
            return redirect(url_for("admin_programs"))

        audience_string = ",".join(selected_audience)

        # ================= DATE =================
        parsed_event_date = datetime.strptime(
            request.form["event_date"],
            "%Y-%m-%d"
        ).date()

        parsed_start_time = datetime.strptime(
            request.form["start_time"],
            "%H:%M"
        ).time()

        parsed_end_time = datetime.strptime(
            request.form["end_time"],
            "%H:%M"
        ).time()

        # ================= SAVE PROGRAM =================
        new_program = Program(
            name=request.form["name"],
            description=request.form["description"],
            target_audience=audience_string,
            location=request.form["location"],
            event_date=parsed_event_date,
            start_time=parsed_start_time,
            end_time=parsed_end_time,
            category="awareness",
            image=filename
        )

        db.session.add(new_program)
        db.session.commit()
        log_action(
    section="program",
    action="create",
    target_type="program",
    target_id=new_program.id,
    description=f"Created program: {new_program.name}"
)


        flash("Program added successfully.", "success")
        return redirect(url_for("admin_programs"))

    # ================= FILTER =================

    now = datetime.now()
    today = now.date()
    current_time = now.time()

    if view == "past":
        programs = Program.query.filter(
            (Program.event_date < today) |
            (
                (Program.event_date == today) &
                (Program.end_time < current_time)
            )
        ).order_by(Program.event_date.desc()).all()
    else:
        programs = Program.query.filter(
            (Program.event_date > today) |
            (
                (Program.event_date == today) &
                (Program.end_time >= current_time)
            )
        ).order_by(Program.event_date.asc()).all()

    return render_template(
        "admin_programs.html",
        programs=programs,
        view=view
    ) 










# =========================
# EDIT PROGRAM
# =========================
from werkzeug.utils import secure_filename
import os
from datetime import datetime

@app.route("/admin/programs/edit/<int:program_id>", methods=["GET", "POST"])
@admin_required
def edit_program(program_id):

    program = Program.query.get_or_404(program_id)

    if request.method == "POST":

        # ================= BASIC FIELDS =================
        program.name = request.form["name"]
        program.description = request.form["description"]
        program.location = request.form["location"]

        program.event_date = datetime.strptime(
            request.form["event_date"],
            "%Y-%m-%d"
        ).date()

        program.start_time = datetime.strptime(
            request.form["start_time"],
            "%H:%M"
        ).time()

        program.end_time = datetime.strptime(
            request.form["end_time"],
            "%H:%M"
        ).time()

        # ================= TARGET AUDIENCE =================
        selected_audience = request.form.getlist("target_audience")
        other = request.form.get("other_audience")

        if other and other.strip():
            selected_audience.append(other.strip())

        program.target_audience = ",".join(selected_audience)

        upload_folder = os.path.join(
            app.root_path,
            "static",
            "uploads",
            "programs"
        )

        os.makedirs(upload_folder, exist_ok=True)

        # ================= CASE 1: NEW FILE UPLOADED =================
        image_file = request.files.get("image")

        if image_file and image_file.filename != "":
            filename = process_and_save_image(
                image_file,
                upload_folder
            )
            program.image = filename

        # ================= CASE 2: CROPPED IMAGE =================
        cropped_data = request.form.get("cropped_image")

        if cropped_data and program.image:

            existing_path = os.path.join(
                upload_folder,
                program.image
            )

            process_and_overwrite_cropped_image(
                cropped_data,
                existing_path
            )

        db.session.commit()
        log_action(
    section="program",
    action="update",
    target_type="program",
    target_id=program.id,
    description=f"Updated program: {program.name}"
)


        flash("Program updated successfully.", "success")
        return redirect(url_for("admin_programs"))

    return render_template(
        "edit_program.html",
        program=program
    )







@app.route("/admin/programs/delete/<int:program_id>", methods=["POST"])
@admin_required
def delete_program(program_id):

    program = Program.query.get_or_404(program_id)

    import os

    # ================= STORE NAME FOR LOGGING =================
    program_name = program.name

    # ================= DELETE IMAGE FILE =================
    if program.image:
        image_path = os.path.join(
            app.root_path,
            "static",
            "uploads",
            "programs",
            program.image
        )

        if os.path.exists(image_path):
            os.remove(image_path)

    # ================= LOG BEFORE DELETE =================
    log_action(
    section="program",
    action="delete",
    target_type="program",
    target_id=program_id,
    description=f"Deleted program: {program_name}"
)


    # ================= DELETE FROM DATABASE =================
    try:
        db.session.delete(program)
        db.session.commit()
    except:
        db.session.rollback()
        flash("Error deleting program.", "danger")



    flash("Program deleted successfully.", "success")

    return redirect(url_for("admin_programs"))



# =========================
# ADMIN ANIMALS
# =========================
@app.route("/admin/animals", methods=["GET", "POST"])
@admin_required
def admin_animals():

    if request.method == "POST":

        new_animal = Animal(
            common_name=request.form["common_name"],
            scientific_name=request.form["scientific_name"],
            habitat=request.form["habitat"],
            conservation_status=request.form["conservation_status"],
            description=request.form["description"]
        )

        # ================= CLOUDINARY IMAGE HANDLING =================
        image_url = None

        cropped_data = request.form.get("cropped_image")
        image_file = request.files.get("image")

        try:
            # If cropped image (base64 from cropper)
            if cropped_data:
                upload_result = cloudinary.uploader.upload(
                    cropped_data,
                    folder="bbp/animals",
                    public_id=f"animal_{uuid4().hex}"
                )
                image_url = upload_result["secure_url"]

            # If normal file upload
            elif image_file and image_file.filename != "":
                upload_result = cloudinary.uploader.upload(
                    image_file,
                    folder="bbp/animals",
                    public_id=f"animal_{uuid4().hex}"
                )
                image_url = upload_result["secure_url"]

        except Exception as e:
            print("Cloudinary Upload Error:", e)
            flash("Image upload failed.", "danger")
            return redirect(url_for("admin_animals"))

        new_animal.image = image_url
        # ===============================================================

        db.session.add(new_animal)
        db.session.commit()

        log_action(
            section="animal",
            action="create",
            target_type="animal",
            target_id=new_animal.id,
            description=f"Added animal: {new_animal.common_name}"
        )

        flash("Animal added successfully.", "success")
        return redirect(url_for("admin_animals"))

    animals = Animal.query.all()

    return render_template(
        "admin_animals.html",
        animals=animals,
        conservation_statuses=CONSERVATION_STATUSES,
        habitats=HABITATS
    )



@app.route("/gallery")
def gallery():
    images = Gallery.query.order_by(Gallery.id.desc()).all()
    return render_template("gallery.html", images=images)

from utils.image_processor import process_and_save_base64_image

@app.route("/admin/gallery", methods=["GET", "POST"])
@admin_required
def admin_gallery():

    upload_folder = os.path.join(
        app.root_path,
        "static",
        "uploads",
        "gallery"
    )
    os.makedirs(upload_folder, exist_ok=True)

    if request.method == "POST":

        data = request.get_json()
        images = data.get("images", [])

        if not images:
            return {"success": False}

        saved_count = 0

        for base64_data in images:
            try:
                filename = process_and_save_base64_image(
                    base64_data,
                    upload_folder
                )

                new_image = Gallery(image=filename)
                db.session.add(new_image)
                saved_count += 1

            except Exception:
                continue

        if saved_count > 0:
            db.session.commit()

            log_action(
        section="gallery",
        action="create",
        target_type="image",
        description=f"Uploaded {saved_count} gallery image(s)"
    )

            return {"success": True}


        return {"success": False}

    images = Gallery.query.order_by(Gallery.id.desc()).all()
    return render_template("admin_gallery.html", images=images)








@app.route("/admin/gallery/delete/<int:image_id>", methods=["POST"])
@admin_required
def delete_gallery_image(image_id):

    image = Gallery.query.get_or_404(image_id)

    # Delete file from disk
    if image.image:
        image_path = os.path.join(
            app.root_path,
            "static",
            "uploads",
            "gallery",
            image.image
        )

        if os.path.exists(image_path):
            os.remove(image_path)

    # Delete DB record
    db.session.delete(image)
    db.session.commit()
    log_action(
    section="gallery",
    action="delete",
    target_type="image",
    target_id=image.id,
    description=f"Deleted gallery image: {image.image}"
)


    flash("Image deleted successfully.", "success")
    return redirect(url_for("admin_gallery"))






@app.route("/admin/gallery/edit/<int:image_id>", methods=["GET", "POST"])
@admin_required
def edit_gallery_image(image_id):

    image = Gallery.query.get_or_404(image_id)

    # ---------- GET: return image URL ----------
    if request.method == "GET":
        return {
            "image_url": url_for(
                "static",
                filename=f"uploads/gallery/{image.image}"
            )
        }

    # ---------- POST: save cropped image ----------
    data = request.get_json()
    base64_data = data.get("image")

    if base64_data:

        path = os.path.join(
            app.root_path,
            "static",
            "uploads",
            "gallery",
            image.image
        )

        process_and_overwrite_cropped_image(base64_data, path)

        return {"success": True}

    return {"success": False}




@app.route("/admin/animals/edit/<int:animal_id>", methods=["GET", "POST"])
@admin_required
def edit_animal(animal_id):

    animal = Animal.query.get_or_404(animal_id)

    if request.method == "POST":

        # ================= BASIC FIELDS =================
        animal.common_name = request.form["common_name"]
        animal.scientific_name = request.form["scientific_name"]
        animal.habitat = request.form["habitat"]
        animal.conservation_status = request.form["conservation_status"]
        animal.description = request.form["description"]

        cropped_data = request.form.get("cropped_image")
        image_file = request.files.get("image")

        upload_folder = os.path.join(app.static_folder, "images", "animals")

        # ================= CROPPED IMAGE =================
        if cropped_data:
            filename = f"{uuid4().hex}.webp"
            save_path = os.path.join(upload_folder, filename)

            process_and_overwrite_cropped_image(
                cropped_data,
                save_path
            )

            animal.image = filename

        # ================= NORMAL UPLOAD =================
        elif image_file and image_file.filename != "":
            filename = process_and_save_image(
                image_file,
                upload_folder
            )

            animal.image = filename

        db.session.commit()
        log_action(
    section="animal",
    action="update",
    target_type="animal",
    target_id=animal.id,
    description=f"Updated animal: {animal.common_name}"
)


        flash("Animal updated successfully.", "success")
        return redirect(url_for("admin_animals"))

    return render_template(
        "edit_animal.html",
        animal=animal,
        conservation_statuses=CONSERVATION_STATUSES,
        habitats=HABITATS
    )







@app.route("/admin/animals/delete/<int:animal_id>")
@admin_required
def delete_animal(animal_id):
    animal = Animal.query.get_or_404(animal_id)

    # STEP A: store data BEFORE delete
    animal_name = animal.common_name

    # STEP B: log BEFORE delete
    log_action(
    section="animal",
    action="delete",
    target_type="animal",
    target_id=animal_id,
    description=f"Deleted animal: {animal_name}"
)


    # STEP C: delete + commit
    db.session.delete(animal)
    db.session.commit()

    return redirect(url_for("admin_animals"))

import base64
import io
from PIL import Image as PILImage

import uuid


@app.route("/admin/outreach", methods=["GET", "POST"])
@admin_required
def admin_outreach():

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    if request.method == "POST":

        title = request.form.get("title")
        description = request.form.get("description")

        # ================= IMAGE PROCESSING =================
        cropped_base64 = request.form.get("cropped_image")
        image_file = request.files.get("image")

        filename = None

        if cropped_base64 or (image_file and image_file.filename != ""):

            os.makedirs("static/uploads/outreach", exist_ok=True)

            unique_name = f"{uuid.uuid4().hex}.webp"
            save_path = os.path.join("static/uploads/outreach", unique_name)

            try:
                if cropped_base64:
                    image_data = cropped_base64.split(",")[1]
                    image_bytes = base64.b64decode(image_data)
                    image = PILImage.open(io.BytesIO(image_bytes))
                else:
                    image = PILImage.open(image_file)

                image = image.convert("RGB")
                image.save(save_path, "WEBP", quality=85, optimize=True)

                filename = unique_name

            except Exception as e:
                print("IMAGE ERROR:", e)
                flash("Image processing failed.", "danger")
                return redirect(url_for("admin_outreach"))

        # ================= DATABASE SAVE =================
        new_post = Outreach(
            title=title,
            description=description,
            image=filename
        )

        db.session.add(new_post)
        db.session.commit()
        log_action(
    section="outreach",
    action="create",
    target_type="post",
    target_id=new_post.id,
    description=f"Created outreach post: {title}"
)


        flash("Outreach content added successfully!", "success")

        return redirect(url_for("admin_outreach"))

    # ================= FETCH POSTS =================
    posts = Outreach.query.order_by(Outreach.created_at.desc()).all()

    return render_template("admin_outreach.html", posts=posts)






@app.route("/admin/outreach/edit/<int:id>", methods=["POST"])
@admin_required
def edit_outreach(id):

    post = Outreach.query.get_or_404(id)

    post.title = request.form.get("title")
    post.description = request.form.get("description")

    cropped_base64 = request.form.get("cropped_image")
    image_file = request.files.get("image")

    if cropped_base64 or (image_file and image_file.filename != ""):

        os.makedirs("static/uploads/outreach", exist_ok=True)

        unique_name = f"{uuid.uuid4().hex}.webp"
        save_path = os.path.join("static/uploads/outreach", unique_name)

        if cropped_base64:
            image_data = cropped_base64.split(",")[1]
            image_bytes = base64.b64decode(image_data)
            image = PILImage.open(io.BytesIO(image_bytes))
        else:
            image = PILImage.open(image_file)

        image = image.convert("RGB")
        image.save(save_path, "WEBP", quality=85, optimize=True)

        # delete old image
        if post.image:
            old_path = os.path.join("static/uploads/outreach", post.image)
            if os.path.exists(old_path):
                os.remove(old_path)

        post.image = unique_name

    db.session.commit()
    log_action(
    section="outreach",
    action="update",
    target_type="post",
    target_id=post.id,
    description=f"Updated outreach post: {post.title}"
)

    flash("Outreach updated successfully.", "success")

    return redirect(url_for("admin_outreach"))



@app.route("/admin/outreach/delete/<int:id>", methods=["POST"])
def delete_outreach(id):

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    post = Outreach.query.get_or_404(id)

    # Delete image file from folder
    if post.image:
        image_path = os.path.join("static/uploads/outreach", post.image)
        if os.path.exists(image_path):
            os.remove(image_path)

    db.session.delete(post)
    db.session.commit()
    log_action(
    section="outreach",
    action="delete",
    target_type="post",
    target_id=post.id,
    description=f"Deleted outreach post: {post.title}"
)

    flash("Outreach deleted successfully.", "success")

    return redirect(url_for("admin_outreach"))



@app.route("/outreach")
def outreach():

    posts = Outreach.query.order_by(Outreach.created_at.desc()).all()

    return render_template("outreach.html", posts=posts)


@app.route("/outreach/<int:id>")
def outreach_detail(id):

    post = Outreach.query.get_or_404(id)

    # Safe increment
    if post.views is None:
        post.views = 1
    else:
        post.views += 1

    db.session.commit()

    return render_template("outreach_detail.html", post=post)

@app.route("/logout", methods=["POST"])
def logout():
    session.clear()
    flash("Logged out successfully.", "success")
    return redirect(url_for("home"))


from datetime import datetime, timedelta


@app.route("/admin/audit")
@admin_required
def admin_audit():

    page = request.args.get("page", 1, type=int)
    section = request.args.get("section")
    action = request.args.get("action")
    search = request.args.get("search")
    start_date = request.args.get("start_date")
    end_date = request.args.get("end_date")

    query = AuditLog.query

    # ================= FILTERS =================

    if section:
        query = query.filter_by(section=section)

    if action:
        query = query.filter_by(action=action)

    if search:
        query = query.filter(
            AuditLog.description.ilike(f"%{search}%")
        )

    # ================= DATE FILTERING =================

    if start_date:
        try:
            start = datetime.strptime(start_date, "%Y-%m-%d")
            query = query.filter(AuditLog.timestamp >= start)
        except:
            pass

    if end_date:
        try:
            end = datetime.strptime(end_date, "%Y-%m-%d") + timedelta(days=1)
            query = query.filter(AuditLog.timestamp < end)
        except:
            pass

    # ================= PAGINATION =================

    logs = query.order_by(AuditLog.timestamp.desc()) \
                .paginate(page=page, per_page=15)

    # ================= SECTION COUNTS (GLOBAL) =================

    section_counts = db.session.query(
        AuditLog.section,
        func.count(AuditLog.id)
    ).group_by(AuditLog.section).all()

    section_dict = dict(section_counts)

    total_logs = sum(section_dict.values())

    # ================= DISTINCT FILTER OPTIONS =================

    sections = [s[0] for s in db.session.query(AuditLog.section).distinct().all()]
    actions = [a[0] for a in db.session.query(AuditLog.action).distinct().all()]

    from datetime import datetime, timedelta

    section_stats = db.session.query(
    AuditLog.section,
    func.count(AuditLog.id)
    ).group_by(AuditLog.section).all()
    
    action_stats = db.session.query(
    AuditLog.action,
    func.count(AuditLog.id)
    ).group_by(AuditLog.action).all()
    
    seven_days_ago = datetime.utcnow() - timedelta(days=7)

    daily_trend = db.session.query(
    func.date(AuditLog.timestamp),
    func.count(AuditLog.id)
    ).filter(
    AuditLog.timestamp >= seven_days_ago
    ).group_by(
    func.date(AuditLog.timestamp)
    ).all()

    
    return render_template(
    "admin_audit.html",
    logs=logs,
    sections=sections,
    actions=actions,
    section_stats=section_stats,
    section_dict=section_dict,
    action_stats=action_stats,
    daily_trend=daily_trend,
    total_logs=sum(section_dict.values())
)


from datetime import datetime

@app.template_filter('datetimeformat')
def datetimeformat(value):
    try:
        date_obj = datetime.strptime(value, "%Y-%m-%d")
        return date_obj.strftime("%d-%m-%Y")
    except:
        return value


@app.route("/admin/logout", methods=["POST"])
def admin_logout():
    session.pop("admin_logged_in", None)
    flash("Logged out successfully.", "success")
    return redirect(url_for("admin_login"))

import csv
from io import StringIO
from flask import Response

@app.route("/admin/audit/export")
@admin_required
def export_audit_logs():

    logs = AuditLog.query.order_by(AuditLog.timestamp.desc()).all()

    output = StringIO()
    writer = csv.writer(output)

    writer.writerow([
        "Timestamp",
        "Actor ID",
        "Actor Role",
        "Section",
        "Action",
        "Target Type",
        "Target ID",
        "Description",
        "IP Address"
    ])

    for log in logs:
        writer.writerow([
            log.timestamp,
            log.actor_id,
            log.actor_role,
            log.section,
            log.action,
            log.target_type,
            log.target_id,
            log.description,
            log.ip_address
        ])

    output.seek(0)

    return Response(
        output,
        mimetype="text/csv",
        headers={
            "Content-Disposition": "attachment; filename=audit_logs.csv"
        }
    )



@app.route("/about-education")
def about_education():
    return render_template("about_education.html")


from datetime import date

from datetime import datetime

@app.route("/programs_public")
def programs_public():

    now = datetime.now()
    today = now.date()
    current_time = now.time()

    programs = Program.query.filter(
        (Program.event_date > today) |
        (
            (Program.event_date == today) &
            (Program.end_time >= current_time)
        )
    ).order_by(Program.event_date.asc()).all()

    return render_template(
        "programs_public.html",
        programs=programs
    )



from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Image as RLImage,
    Table,
    TableStyle,
    PageBreak
)
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from io import BytesIO
from flask import send_file
import os


# ================= HEADER =================

def add_page_design(c, doc):
    width, height = A4

    logo_path = os.path.join(app.root_path, "static/images/bbp-logo.png")

    if os.path.exists(logo_path):
        logo = ImageReader(logo_path)
        logo_width, logo_height = logo.getSize()

        aspect = logo_height / float(logo_width)
        desired_width = 50
        desired_height = desired_width * aspect

        c.drawImage(
            logo_path,
            40,
            height - 75,
            width=desired_width,
            height=desired_height,
            mask='auto'
        )

    # Title
    c.setFillColor(colors.black)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(110, height - 45, "Bannerughatta Biological Park")

    c.setFont("Helvetica", 10)
    c.drawString(110, height - 60, "Copy of Animal Report in website ")

    # Page number
    c.setFont("Helvetica", 9)
    c.drawRightString(width - 40, 20, f"Page {doc.page}")


# ================= MAIN ROUTE =================

@app.route("/admin/backup/animals/pdf")
@admin_required
def backup_animals_pdf():

    animals = Animal.query.order_by(Animal.id.asc()).all()

    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=100,
        bottomMargin=60
    )

    elements = []
    styles = getSampleStyleSheet()
    normal_style = styles["Normal"]

    # IUCN Color Mapping
    def get_iucn_color(status):
        status = (status or "").lower()

        if "critically" in status:
            return colors.HexColor("#8B0000")
        elif "endangered" in status:
            return colors.red
        elif "vulnerable" in status:
            return colors.orange
        elif "near" in status:
            return colors.HexColor("#FF8C00")
        else:
            return colors.green

    animals_per_page = 2
    count = 0

    for animal in animals:

        content = []

        # Animal Image
        if animal.image:
            image_path = os.path.join(
                app.root_path,
                "static/images/animals",
                animal.image
            )
            if os.path.exists(image_path):
                img = RLImage(image_path, width=3*inch, height=2*inch)
                content.append(img)
                content.append(Spacer(1, 0.2 * inch))

        # ID Badge
        badge_style = ParagraphStyle(
            name='Badge',
            backColor=colors.HexColor("#E8F5E9"),
            textColor=colors.HexColor("#1B5E20"),
            fontSize=9,
            leftIndent=4,
            rightIndent=4
        )

        content.append(Paragraph(f"<b>ID #{animal.id}</b>", badge_style))
        content.append(Spacer(1, 0.1 * inch))

        # Name
        content.append(Paragraph(
            f"<b>{animal.common_name}</b>",
            styles["Heading3"]
        ))

        content.append(Paragraph(
            f"<i>{animal.scientific_name}</i>",
            normal_style
        ))

        content.append(Spacer(1, 0.1 * inch))

        # Habitat
        content.append(Paragraph(
            f"<b>Habitat:</b> {animal.habitat}",
            normal_style
        ))

        # Description
        if animal.description:
            content.append(Spacer(1, 0.1 * inch))
            content.append(Paragraph(
                animal.description,
                normal_style
            ))

        # IUCN Status
        iucn_color = get_iucn_color(animal.conservation_status)

        content.append(Spacer(1, 0.15 * inch))
        content.append(Paragraph(
            f"<font color='{iucn_color.hexval()}'><b>IUCN: {animal.conservation_status}</b></font>",
            normal_style
        ))

        # Wrap each animal inside bordered card
        card = Table([[content]], colWidths=[6.7 * inch])
        card.setStyle(TableStyle([
            ("BOX", (0,0), (-1,-1), 0.5, colors.grey),
            ("LEFTPADDING", (0,0), (-1,-1), 12),
            ("RIGHTPADDING", (0,0), (-1,-1), 12),
            ("TOPPADDING", (0,0), (-1,-1), 12),
            ("BOTTOMPADDING", (0,0), (-1,-1), 12),
        ]))

        elements.append(card)
        elements.append(Spacer(1, 0.4 * inch))

        count += 1

        # Page break after 2 animals
        if count % animals_per_page == 0:
            elements.append(PageBreak())

    doc.build(
        elements,
        onFirstPage=add_page_design,
        onLaterPages=add_page_design
    )

    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="BBP_Animals_Official_Report.pdf",
        mimetype="application/pdf"
    )



from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import send_file
from io import BytesIO
from PIL import Image
import tempfile
import os


@app.route("/admin/backup/animals/docx")
@admin_required
def backup_animals_docx():

    animals = Animal.query.order_by(Animal.id.asc()).all()

    doc = Document()

    # ===== TITLE =====
    doc.add_heading("Bannerghatta Biological Park", level=1)
    doc.add_paragraph("Official Animal Report")

    for index, animal in enumerate(animals):

        # Animal Name
        doc.add_heading(animal.common_name, level=2)

        # Scientific Name (italic)
        sci_para = doc.add_paragraph()
        sci_run = sci_para.add_run(animal.scientific_name)
        sci_run.italic = True

        # Habitat
        doc.add_paragraph(f"Habitat: {animal.habitat}")

        # Description
        if animal.description:
            doc.add_paragraph(animal.description)

        # IUCN
        doc.add_paragraph(f"IUCN Status: {animal.conservation_status}")

        # ===== IMAGE HANDLING (WEBP SAFE) =====
        if animal.image:
            image_path = os.path.join(
                app.root_path,
                "static/images/animals",
                animal.image
            )

            if os.path.exists(image_path):
                try:
                    # Open WEBP
                    img = Image.open(image_path).convert("RGB")

                    # Create temporary PNG file
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                        temp_path = tmp.name
                        img.save(temp_path, "PNG")

                    # Add picture to document
                    doc.add_picture(temp_path, width=Inches(3))

                    # Center the image
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Remove temporary file
                    os.remove(temp_path)

                except Exception as e:
                    print("Image conversion failed:", e)

        # Page break except last
        if index < len(animals) - 1:
            doc.add_page_break()

    # ===== SAVE TO MEMORY =====
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="BBP_Animals_Official_Report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )














@app.route("/admin/backup/programs")
def backup_programs():

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    programs = Program.query.all()

    def generate():
        yield "Name,Target Audience,Duration,Description\n"
        for p in programs:
            yield f"{p.name},{p.target_audience},{p.duration},{p.description}\n"

    return Response(
        generate(),
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment;filename=programs_backup.csv"}
    )
@app.route("/admin/backup/events")
def backup_events():

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    events = Event.query.all()

    def generate():
        yield "Title,Date,Description\n"
        for e in events:
            yield f"{e.title},{e.event_date},{e.description}\n"

    return Response(
        generate(),
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment;filename=events_backup.csv"}
    )

@app.route("/admin/backup/gallery")
def backup_gallery():

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    gallery_path = "static/images/gallery"

    if not os.path.exists(gallery_path):
        return "Gallery folder not found", 404

    memory_file = io.BytesIO()

    with zipfile.ZipFile(memory_file, "w", zipfile.ZIP_DEFLATED) as zf:
        for filename in os.listdir(gallery_path):
            file_path = os.path.join(gallery_path, filename)
            if os.path.isfile(file_path):
                zf.write(file_path, arcname=filename)

    memory_file.seek(0)

    return Response(
        memory_file,
        mimetype="application/zip",
        headers={
            "Content-Disposition": "attachment; filename=gallery_backup.zip"
        }
    )
import os
from flask import send_file
from datetime import datetime


@app.route("/admin/backup/full")
@admin_required
def backup_full_system():

    db_path = os.path.join(BASE_DIR, "site.db")

    if not os.path.exists(db_path):
        flash("Database file not found!", "danger")
        return redirect(url_for("admin_backup"))

    # Create timestamp filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_filename = f"site_backup_{timestamp}.db"
    backup_path = os.path.join("backups", backup_filename)

    shutil.copy(db_path, backup_path)
    # Auto delete old backups (keep latest 10)
    backup_folder = "backups"
    files = sorted(
        os.listdir(backup_folder),
        reverse=True
)

    if len(files) > 10:
        for old_file in files[10:]:
            os.remove(os.path.join(backup_folder, old_file))

    flash("Full system backup created successfully!", "success")

    return send_file(
        backup_path,
        as_attachment=True
    )
    
    
@app.route("/admin/backup/history")
@admin_required
def backup_history():

    backup_folder = "backups"

    if not os.path.exists(backup_folder):
        os.makedirs(backup_folder)

    files = os.listdir(backup_folder)
    files.sort(reverse=True)

    return render_template(
        "backup_history.html",
        files=files
    )



@app.route("/admin/backup/download/<filename>")
@admin_required
def download_backup(filename):

    path = os.path.join("backups", filename)

    if not os.path.exists(path):
        flash("File not found!", "danger")
        return redirect(url_for("backup_history"))

    return send_file(path, as_attachment=True)

@app.route("/admin/backup/restore", methods=["POST"])
@admin_required
def restore_backup():

    file = request.files.get("backup_file")

    if not file:
        flash("No file selected!", "danger")
        return redirect(url_for("backup_history"))

    db_path = os.path.join(BASE_DIR, "site.db")

    file.save(db_path)

    flash("System restored successfully. Restart server.", "success")

    return redirect(url_for("admin_dashboard"))

# =========================
# VOLUNTEER AUTH ROUTES
# =========================

@app.route("/volunteer/register", methods=["GET", "POST"])
def volunteer_register():

    if request.method == "POST":

        name = request.form.get("name")
        email = request.form.get("email")
        password = request.form.get("password")

        existing = User.query.filter_by(email=email).first()

        # 🔹 CASE 1 — Already exists
        if existing:

            if existing.status == "pending":
                flash("Your application is already under review.", "warning")
                return redirect(url_for("volunteer_register"))

            if existing.status == "approved":
                flash("You are already approved. Please login.", "info")
                return redirect(url_for("volunteer_login"))

            if existing.status == "rejected":
                # 🔥 Reset application
                existing.name = name
                existing.password_hash = generate_password_hash(password, method="pbkdf2:sha256")
                existing.status = "pending"

                db.session.commit()

                flash("Application resubmitted successfully.", "success")
                return redirect(url_for("volunteer_login"))

        # 🔹 CASE 2 — New user
        hashed_pw = generate_password_hash(password, method="pbkdf2:sha256")


        new_user = User(
            name=name,
            email=email,
            password_hash=hashed_pw,
            role="volunteer",
            status="pending"
        )

        db.session.add(new_user)
        db.session.commit()

        flash("Registration successful. Awaiting approval.", "success")
        return redirect(url_for("volunteer_login"))

    return render_template("volunteer_register.html")


from werkzeug.utils import secure_filename
from datetime import datetime
import os
import base64
from PIL import Image
from io import BytesIO


@app.route("/intern/register", methods=["GET", "POST"])
def intern_register():

    if request.method == "POST":
        print("FORM SUBMITTED")

        # ================= BASIC INFO =================
        name = request.form.get("name")
        email = request.form.get("email", "").strip()
        password = request.form.get("password")

        course = request.form.get("course")
        year = request.form.get("year")
        semester = request.form.get("semester")
        college = request.form.get("college")

        phone = request.form.get("phone", "").strip()

        # ================= PHONE VALIDATION =================
        if not phone.isdigit():
            flash("Phone number must contain digits only.", "danger")
            return redirect(request.url)

        if len(phone) != 10 or phone[0] not in "6789":
            flash("Enter a valid 10-digit Indian mobile number.", "danger")
            return redirect(request.url)

        full_phone = "+91" + phone

        objective = request.form.get("objective")
        start_date = request.form.get("start_date")
        end_date = request.form.get("end_date")

        # ================= SUBJECT VALIDATION =================
        selected_subjects = request.form.getlist("internship_subjects")

        if not selected_subjects:
            flash("Please select at least one internship subject.", "danger")
            return redirect(request.url)

        subjects_string = ",".join(selected_subjects)

        # ================= DUPLICATE EMAIL CHECK =================
        existing_user = User.query.filter_by(email=email).first()

        if existing_user:

            if existing_user.status != "rejected":
                flash("Email already registered.", "danger")
                return redirect(url_for("intern_register"))

            if existing_user.intern_profile:
                db.session.delete(existing_user.intern_profile)

            db.session.delete(existing_user)
            db.session.commit()

        # ================= PHOTO (CROPPED BASE64) =================
        import base64
        from PIL import Image
        from io import BytesIO
        cropped_image = request.form.get("cropped_image")

        if not cropped_image:
            flash("Please crop your photo before submitting.", "danger")
            return redirect(request.url)

        try:
            # Split base64
            header, encoded = cropped_image.split(",", 1)
            image_data = base64.b64decode(encoded)

            # Open image
            img = Image.open(BytesIO(image_data))
            img = img.convert("RGB")

             # Convert to WEBP
            buffer = BytesIO()
            img.save(buffer, format="WEBP", quality=80)
            buffer.seek(0)

            # Upload to Cloudinary
            upload_result = cloudinary.uploader.upload(
            buffer,
            folder="bbp/interns",
            resource_type="image"
            )

            # Get secure URL
            photo_url = upload_result["secure_url"]

        except Exception as e:
            print("CLOUDINARY ERROR:", e)
            flash("Error processing image.", "danger")
            return redirect(request.url)

        # ================= CREATE USER =================
        password_hash = generate_password_hash(
            password, method="pbkdf2:sha256"
        )

        user = User(
            name=name,
            email=email,
            password_hash=password_hash,
            role="intern",
            status="pending"
        )

        db.session.add(user)
        db.session.commit()

        # ================= CREATE PROFILE =================
        profile = InternProfile(
            user_id=user.id,
            course=course,
            year=year,
            semester=semester,
            college=college,
            phone=full_phone,
            start_date=datetime.strptime(start_date, "%Y-%m-%d"),
            end_date=datetime.strptime(end_date, "%Y-%m-%d"),
            objective=objective,
            internship_subjects=subjects_string,
            application_status="submitted",
            photo=photo_url
        )

        db.session.add(profile)
        db.session.commit()

        flash("Application submitted. Awaiting approval.", "success")
        return redirect(url_for("intern_login"))

    return render_template("intern_register.html")






@app.route("/forgot-password/<role>", methods=["GET", "POST"])
def forgot_password(role):
    if request.method == "POST":
        email = request.form.get("email")

        user = User.query.filter_by(email=email, role=role).first()

        if user:
            token = str(uuid.uuid4())
            user.reset_token = token
            db.session.commit()

            reset_url = url_for(
                "reset_password",
                role=role,
                token=token,
                _external=True
            )

            print("Reset Link:", reset_url)

            flash("Reset link generated. Check console.")

        else:
            flash("User not found.")

    return render_template("forgot_password.html", role=role)




@app.route("/reset-password/<role>/<token>", methods=["GET", "POST"])
def reset_password(role, token):

    user = User.query.filter_by(reset_token=token, role=role).first()

    if not user:
        flash("Invalid or expired token.")
        return redirect(url_for("forgot_password", role=role))

    if request.method == "POST":
        new_password = request.form.get("password")

        user.password_hash = generate_password_hash(
    new_password,
    method="pbkdf2:sha256"
)
        user.reset_token = None
        db.session.commit()

        flash("Password reset successful.")
        return redirect(url_for(f"{role}_login"))

    return render_template("reset_password.html", role=role)










@app.route("/volunteer/dashboard")
def volunteer_dashboard():

    print("VOLUNTEER DASHBOARD HIT")
    print("SESSION:", dict(session))

    if not session.get("volunteer_id"):
        print("NO VOLUNTEER SESSION — REDIRECTING")
        return redirect(url_for("volunteer_login"))

    return render_template("volunteer_dashboard.html")





@app.route("/intern/login", methods=["GET", "POST"])
def intern_login():

    if request.method == "POST":

        email = request.form.get("email")
        password = request.form.get("password")

        user = User.query.filter_by(email=email, role="intern").first()

        if not user:
            flash("Account not found", "danger")
            return redirect(url_for("intern_login"))

        if not check_password_hash(user.password_hash, password):
            flash("Invalid password", "danger")
            return redirect(url_for("intern_login"))

        # STATUS-BASED APPROVAL CHECK
        if user.status == "pending":
            flash("Application submitted. Awaiting admin approval.", "warning")
            return redirect(url_for("intern_login"))

        if user.status == "rejected":
            flash("Your application was rejected.", "danger")
            return redirect(url_for("intern_login"))

        if user.status != "approved":
            flash("Account not approved yet.", "warning")
            return redirect(url_for("intern_login"))

        # If approved → allow login
        session["user_id"] = user.id
        session["user_name"] = user.name
        session["role"] = "intern"
        session.permanent = True

        return redirect(url_for("intern_dashboard"))

    return render_template("intern_login.html")










from werkzeug.utils import secure_filename
import os
from datetime import datetime, date


from datetime import datetime, date

from datetime import datetime, date

from datetime import date
from datetime import date
from flask import session, redirect, url_for, render_template



from datetime import date
from flask import session, redirect, url_for, render_template


@app.route("/intern/dashboard")
def intern_dashboard():

    # 🔐 Authentication check
    if "user_id" not in session or session.get("role") != "intern":
        return redirect(url_for("intern_login"))

    user = User.query.get(session["user_id"])
    if not user:
        return redirect(url_for("intern_login"))

    profile = user.intern_profile
    today = date.today()

    # ================= TODAY LOG =================
    today_log = InternDailyLog.query.filter_by(
        intern_id=user.id,
        date=today
    ).first()

    # ================= PREVIOUS LOGS =================
    previous_logs = InternDailyLog.query.filter(
        InternDailyLog.intern_id == user.id,
        InternDailyLog.date < today
    ).order_by(InternDailyLog.date.desc()).all()

    # ================= DAILY LOGIN CONTROL =================
    allow_today_login = False

    if profile and not profile.is_completed:

        # Case 1: Admin manually disabled daily logging
        if not profile.daily_log_enabled:
            allow_today_login = False

        # Case 2: Report submitted — disable from next day
        elif profile.report_submission_date:
            if today > profile.report_submission_date:
                allow_today_login = False
            else:
                allow_today_login = True

        # Case 3: Normal active intern
        else:
            allow_today_login = True

    # If internship completed → fully locked
    else:
        allow_today_login = False

    return render_template(
        "intern_dashboard.html",
        user=user,
        profile=profile,
        today_log=today_log,
        previous_logs=previous_logs,
        allow_today_login=allow_today_login
    )








@app.route("/intern/dashboard/day/<int:log_id>")
def intern_day_view(log_id):

    if "user_id" not in session or session.get("role") != "intern":
        return redirect(url_for("intern_login"))

    log = InternDailyLog.query.get_or_404(log_id)

    # security check
    if log.intern_id != session["user_id"]:
        return redirect(url_for("intern_dashboard"))

    return render_template(
        "intern_day_view.html",
        log=log
    )


from datetime import date, datetime
import pytz

@app.route("/intern/today", methods=["GET", "POST"])
def intern_today():

    # 🔐 Session check
    if "user_id" not in session or session.get("role") != "intern":
        return redirect(url_for("intern_login"))

    user = User.query.get(session["user_id"])
    if not user:
        return redirect(url_for("intern_login"))

    profile = user.intern_profile
    today = date.today()

    # 🇮🇳 India Timezone
    india = pytz.timezone("Asia/Kolkata")
    now_ist = datetime.now(india)

    # 🔒 Block conditions
    if not profile:
        return redirect(url_for("intern_dashboard"))

    if profile.is_completed:
        return redirect(url_for("intern_dashboard"))

    if not profile.daily_log_enabled:
        return redirect(url_for("intern_dashboard"))

    if profile.report_submission_date and today > profile.report_submission_date:
        return redirect(url_for("intern_dashboard"))

    # ================= GET OR CREATE TODAY LOG =================
    log = InternDailyLog.query.filter_by(
        intern_id=user.id,
        date=today
    ).first()

    if not log:
        log = InternDailyLog(
            intern_id=user.id,
            date=today,
            login_time=now_ist,
            submitted=False
        )
        db.session.add(log)
        db.session.commit()

    # ================= HANDLE FORM SUBMISSION =================
    if request.method == "POST":

        # Always save text fields
        log.summary = request.form.get("summary")
        log.location = request.form.get("location")
        log.schedule = request.form.get("schedule")

        # Submit & Close Day button
        if "submit_today" in request.form:
            log.logout_time = now_ist
            log.submitted = True

        # Logout button (if shown later)
        if "logout_today" in request.form:
            log.logout_time = now_ist

        db.session.commit()

        return redirect(url_for("intern_dashboard"))

    return render_template("intern_today.html", log=log)






@app.route("/admin/intern/<int:intern_id>/enable-logs", methods=["POST"])
def enable_daily_logs(intern_id):

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    profile = InternProfile.query.filter_by(user_id=intern_id).first_or_404()

    profile.daily_log_enabled = True
    db.session.commit()

    flash("Daily logging re-enabled successfully.", "success")
    return redirect(url_for("admin_intern_days", intern_id=intern_id))


@app.route("/intern/login_today", methods=["POST"])
def intern_login_today():

    if "user_id" not in session:
        return redirect(url_for("intern_login"))

    user_id = session["user_id"]
    user = User.query.get_or_404(user_id)
    profile = user.intern_profile

    if not profile.daily_log_enabled:
        flash("Daily logging has been disabled after report submission.", "warning")
        return redirect(url_for("intern_dashboard"))

    if user.intern_profile and user.intern_profile.is_completed:
        flash("Your internship has already been completed.", "warning")
        return redirect(url_for("intern_dashboard"))

    today = date.today()

    existing = InternDailyLog.query.filter_by(
        intern_id=user_id,
        date=today
    ).first()

    if existing:
        return redirect(url_for("intern_today"))

    new_log = InternDailyLog(
        intern_id=user_id,
        date=today,
        login_time=datetime.now()
    )

    db.session.add(new_log)
    db.session.commit()

    return redirect(url_for("intern_today"))




@app.route("/intern/logout", methods=["POST"])
def intern_logout():

    session.pop("intern_logged_in", None)
    session.pop("intern_id", None)

    flash("Logged out successfully.", "success")
    return redirect(url_for("intern_login"))





@app.route("/admin/backup")
@admin_required
def admin_backup():
    return render_template("admin_backup.html")


@app.route("/admin/requests")
@admin_required
def admin_requests():

    volunteer_requests = User.query.filter_by(
        role="volunteer",
        status="pending"
    ).all()

    intern_requests = User.query.filter_by(
        role="intern",
        status="pending"
    ).all()

    return render_template(
        "admin_requests.html",
        volunteer_requests=volunteer_requests,
        intern_requests=intern_requests
    )

    
@app.route("/approve-user/<int:user_id>")
@admin_required
def approve_user(user_id):

    user = User.query.get_or_404(user_id)

    try:
        # Update User table
        user.status = "approved"

        # Update InternProfile table
        if user.role == "intern" and user.intern_profile:
            user.intern_profile.application_status = "approved"
            user.intern_profile.daily_log_enabled = True

        db.session.commit()
        log_action(
    section="intern" if user.role == "intern" else "volunteer",
    action="approve",
    target_type="user",
    target_id=user.id,
    description=f"Approved {user.role}: {user.name}"
)

        flash("User approved successfully.", "success")

    except Exception:
        db.session.rollback()
        flash("Error approving user.", "danger")

    return redirect(url_for("admin_requests"))






import os
from sqlalchemy.exc import SQLAlchemyError

@app.route("/reject-user/<int:user_id>")
@admin_required
def reject_user(user_id):

    user = User.query.get_or_404(user_id)

    try:
        # If intern → remove related data
        if user.role == "intern":

            # Delete logs
            InternDailyLog.query.filter_by(
                intern_id=user.id
            ).delete(synchronize_session=False)

            profile = user.intern_profile

            if profile:

                # Delete certificate file
                if profile.certificate_file:
                    certificate_path = os.path.join(
                        "storage/certificates",
                        profile.certificate_file
                    )
                    if os.path.exists(certificate_path):
                        os.remove(certificate_path)

                # Delete report file
                if profile.report_file:
                    report_path = os.path.join(
                        "static/reports",
                        profile.report_file
                    )
                    if os.path.exists(report_path):
                        os.remove(report_path)

                db.session.delete(profile)

        # Delete user itself
        db.session.delete(user)

        db.session.commit()

        flash("User rejected and permanently removed.", "success")

    except Exception:
        db.session.rollback()
        flash("Error rejecting user.", "danger")

    return redirect(url_for("admin_requests"))



@app.route("/admin/members")
@superadmin_required
def admin_members():
    return render_template("admin_members.html")

@app.route("/admin/members/volunteers")
@superadmin_required
def view_volunteers():
    volunteers = User.query.filter_by(
        role="volunteer",
        status="approved"
    ).all()
    return render_template("admin_volunteers.html", users=volunteers)
@app.route("/admin/members/interns")
@superadmin_required
def view_interns():
    interns = User.query.filter_by(
        role="intern",
        status="approved"
    ).all()
    return render_template("admin_interns.html", users=interns)

import os
from flask import flash, redirect, url_for
from sqlalchemy.exc import SQLAlchemyError


@app.route("/admin/delete-intern/<int:intern_id>", methods=["POST"])
@admin_required
def delete_intern(intern_id):

    intern = User.query.get_or_404(intern_id)

    # 🔐 Safety check — ensure only interns are deleted
    if intern.role != "intern":
        flash("Invalid user type.", "danger")
        return redirect(url_for("admin_interns"))

    try:
        profile = intern.intern_profile

        # ================= DELETE DAILY LOGS =================
        InternDailyLog.query.filter_by(intern_id=intern.id).delete(
            synchronize_session=False
        )

        # ================= DELETE CERTIFICATE FILE =================
        if profile and profile.certificate_file:
            certificate_path = os.path.join(
                "storage/certificates",
                profile.certificate_file
            )
            if os.path.exists(certificate_path):
                os.remove(certificate_path)

        # ================= DELETE REPORT FILE =================
        if profile and profile.report_file:
            report_path = os.path.join(
                "static/reports",
                profile.report_file
            )
            if os.path.exists(report_path):
                os.remove(report_path)

        # ================= DELETE PROFILE =================
        if profile:
            db.session.delete(profile)

        # ================= DELETE USER =================
        db.session.delete(intern)

        db.session.commit()

        flash("Intern and all related data permanently deleted.", "success")

    except SQLAlchemyError:
        db.session.rollback()
        flash("Error deleting intern. Please try again.", "danger")

    return redirect(url_for("admin_intern_records"))





@app.route("/admin/volunteer/delete/<int:user_id>")
@admin_required
def delete_volunteer(user_id):

    user = User.query.get_or_404(user_id)

    if user.role != "volunteer":
        flash("Invalid action.", "danger")
        return redirect(url_for("admin_volunteers"))

    db.session.delete(user)
    db.session.commit()

    flash("Volunteer removed successfully.", "success")
    return redirect(url_for("admin_volunteers"))








@app.errorhandler(403)
def forbidden(e):
    return render_template("403.html"), 403

import csv
from flask import Response


@app.route("/admin/backup/volunteers")
@admin_required
def backup_volunteers():

    volunteers = User.query.filter_by(
        role="volunteer",
        status="approved"
    ).all()

    def generate():
        data = csv.writer(open("volunteers_temp.csv", "w", newline=""))
        yield "Name,Email,Role,Status\n"
        for v in volunteers:
            yield f"{v.name},{v.email},{v.role},{v.status}\n"

    return Response(
        generate(),
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment;filename=volunteers_backup.csv"}
    )
import csv
from io import StringIO
from flask import Response

@app.route("/admin/backups/interns")
@admin_required
def backup_interns():

    output = StringIO()
    writer = csv.writer(output)

    writer.writerow([
        "Intern ID",
        "Name",
        "Email",
        "Status",
        "Log Date",
        "Daily Summary"
    ])

    interns = User.query.filter_by(role="intern").all()

    for intern in interns:
        logs = InternDailyLog.query.filter_by(intern_id=intern.id).all()

        if logs:
            for log in logs:
                writer.writerow([
                    intern.id,
                    intern.name,
                    intern.email,
                    intern.status,
                    log.date,
                    log.summary
                ])
        else:
            writer.writerow([
                intern.id,
                intern.name,
                intern.email,
                intern.status,
                "",
                "No logs submitted"
            ])

    output.seek(0)

    return Response(
        output,
        mimetype="text/csv",
        headers={
            "Content-Disposition": "attachment;filename=intern_backup.csv"
        }
    )


from sqlalchemy import or_

@app.route("/admin/intern-records")
def admin_intern_records():

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    search = request.args.get("search", "").strip()
    selected_subjects = request.args.getlist("subjects")
    page = request.args.get("page", 1, type=int)

    subject_options = [
        "Animal Behaviour",
        "Flora",
        "Butterfly Management",
        "Marketing",
        "Designing"
    ]

    # ================= BASE QUERY =================
    query = User.query.join(InternProfile).filter(
        User.role == "intern",
        InternProfile.application_status == "approved"
    )

    # 🔎 SEARCH FILTER
    if search:
        query = query.filter(
            or_(
                User.name.ilike(f"%{search}%"),
                User.email.ilike(f"%{search}%"),
                InternProfile.college.ilike(f"%{search}%"),
                InternProfile.course.ilike(f"%{search}%")
            )
        )

    # 🎯 SUBJECT FILTER
    if selected_subjects:
        subject_filters = [
            InternProfile.internship_subjects.ilike(f"%{sub}%")
            for sub in selected_subjects
        ]
        query = query.filter(or_(*subject_filters))

    # ================= GLOBAL COUNTS =================
    base_query = User.query.join(InternProfile).filter(
        User.role == "intern",
        InternProfile.application_status == "approved"
    )

    total_count = base_query.count()

    active_count = base_query.filter(
        InternProfile.is_completed == False
    ).count()

    completed_count = base_query.filter(
        InternProfile.is_completed == True
    ).count()

    rejected_count = 0  # Rejected not part of this page anymore

    # 📄 PAGINATION
    pagination = query.order_by(User.id.desc()).paginate(page=page, per_page=6)
    interns = pagination.items

    return render_template(
        "admin_intern_records.html",
        interns=interns,
        pagination=pagination,
        total_count=total_count,
        active_count=active_count,
        completed_count=completed_count,
        rejected_count=rejected_count,
        subject_options=subject_options,
        selected_subjects=selected_subjects,
        search=search
    )





@app.context_processor
def inject_certificate_requests():
    certificate_requests = InternProfile.query.filter_by(
        certificate_requested=True,
        is_completed=False
    ).count()
    return dict(certificate_requests=certificate_requests)



@app.route("/admin/intern/<int:intern_id>/days")
def admin_intern_days(intern_id):

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    intern = User.query.get_or_404(intern_id)

    # Fetch intern profile
    profile = intern.intern_profile

    # If application rejected → show special view
    if profile and profile.application_status == "rejected":
        return render_template(
            "admin_intern_days.html",
            intern=intern,
            logs=[],
            rejected=True
        )

    # Normal case → fetch logs
    logs = InternDailyLog.query.filter_by(
        intern_id=intern.id
    ).order_by(InternDailyLog.date.asc()).all()

    return render_template(
        "admin_intern_days.html",
        intern=intern,
        logs=logs,
        rejected=False
    )

@app.route("/admin/intern/<int:intern_id>/day/<int:log_id>")
def admin_intern_day_detail(intern_id, log_id):

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    intern = User.query.get_or_404(intern_id)

    log = InternDailyLog.query.filter_by(
        id=log_id,
        intern_id=intern.id
    ).first_or_404()

    return render_template(
        "admin_intern_day_detail.html",
        intern=intern,
        log=log
    )



@app.route("/admin/intern/<int:user_id>/day/<int:log_id>")
def admin_day_detail(user_id, log_id):

    intern = User.query.get_or_404(user_id)

    log = InternDailyLog.query.get_or_404(log_id)

    return render_template(
        "admin_day_detail.html",
        intern=intern,
        log=log
    )
@app.route("/admin/intern/day/reset/<int:log_id>", methods=["POST"])
def admin_reset_intern_day(log_id):

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    log = InternDailyLog.query.get_or_404(log_id)

    # Archive only if current data exists
    if log.summary or log.location or log.schedule:
        log.previous_summary = log.summary
        log.previous_location = log.location
        log.previous_schedule = log.schedule

    # Clear current
    log.summary = None
    log.location = None
    log.schedule = None
    log.submitted = False
    log.logout_time = None

    db.session.commit()

    flash("Submission reset. Previous version archived.", "info")

    return redirect(url_for(
        "admin_intern_day_detail",
        intern_id=log.intern_id,
        log_id=log.id
    ))





@app.route("/admin/intern/<int:intern_id>/enable-report", methods=["POST"])
def admin_enable_report(intern_id):

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    intern = User.query.get_or_404(intern_id)

    profile = InternProfile.query.filter_by(user_id=intern.id).first()

    if not profile:
        flash("Intern profile not found. This intern must re-register.", "danger")
        return redirect(url_for("admin_intern_records"))

    profile.report_upload_enabled = True
    db.session.commit()

    flash("Report upload enabled.", "success")
    return redirect(url_for("admin_intern_days", intern_id=intern.id))





import os
from datetime import date
from werkzeug.utils import secure_filename


UPLOAD_FOLDER = "static/reports"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER


@app.route("/intern/upload-report", methods=["POST"])
def upload_intern_report():

    # 🔐 Authentication check
    if "user_id" not in session or session.get("role") != "intern":
        return redirect(url_for("intern_login"))

    user_id = session["user_id"]
    profile = InternProfile.query.filter_by(user_id=user_id).first()

    if not profile:
        return redirect(url_for("intern_dashboard"))

    # 🚫 If upload not enabled
    if not profile.report_upload_enabled:
        return redirect(url_for("intern_dashboard"))

    file = request.files.get("report")

    if not file or not file.filename.endswith(".pdf"):
        return redirect(url_for("intern_dashboard"))

    # 📁 Secure file naming
    filename = secure_filename(f"intern_{user_id}.pdf")
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)

    file.save(filepath)

    # 🕒 Store submission date (TIMELINE CONTROL)
    profile.report_submission_date = date.today()

    # 📄 Store file name
    profile.report_file = filename

    # 🔒 Optional: disable further daily logs immediately (if you still want this)
    # profile.daily_log_enabled = False

    db.session.commit()

    return redirect(url_for("intern_dashboard"))


@app.route("/admin/download-report/<int:intern_id>")
def admin_download_report(intern_id):

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    profile = InternProfile.query.filter_by(user_id=intern_id).first()

    if profile.report_file:
        return redirect(url_for('static', filename=f'reports/{profile.report_file}'))

    return redirect(url_for("admin_intern_days", intern_id=intern_id))



# =========================
# DATABASE INITIALIZATION
# =========================


    # 🔐 Ensure superadmin always exists
    if Admin.query.count() == 0:
        default_password = bcrypt.generate_password_hash("admin123").decode("utf-8")

        super_admin = Admin(
            username="education_officer",
            password_hash=default_password,
            role="superadmin"
        )

        db.session.add(super_admin)
        db.session.commit()

        print("Default superadmin created.")
        
@app.route("/admin/intern/<int:intern_id>/view")
@admin_required
def admin_view_intern_application(intern_id):

    intern = User.query.filter_by(id=intern_id, role="intern").first()

    if not intern:
        flash("Intern not found.", "danger")
        return redirect(url_for("admin_requests"))

    profile = InternProfile.query.filter_by(user_id=intern.id).first()

    if not profile:
        flash("Application profile missing.", "danger")
        return redirect(url_for("admin_requests"))

    # ✅ Auto change to under_review
    if profile.application_status == "submitted":
        profile.application_status = "under_review"
        db.session.commit()

    return render_template(
        "admin_view_intern_application.html",
        intern=intern,
        profile=profile
    )

@app.route("/admin/intern/<int:intern_id>/process", methods=["POST"])
@admin_required
def admin_process_application(intern_id):

    intern = User.query.filter_by(id=intern_id, role="intern").first()

    if not intern:
        flash("Intern not found.", "danger")
        return redirect(url_for("admin_requests"))

    profile = InternProfile.query.filter_by(user_id=intern.id).first()

    if not profile:
        flash("Application profile missing.", "danger")
        return redirect(url_for("admin_requests"))

    action = request.form.get("action")

    if action == "approve":
        intern.status = "approved"
        profile.application_status = "approved"
        flash("Application approved successfully.", "success")

    elif action == "reject":
        intern.status = "rejected"
        profile.application_status = "rejected"
        flash("Application rejected.", "danger")

    db.session.commit()
    log_action(
    section="intern",
    action=action,
    target_type="user",
    target_id=intern.id,
    description=f"{action.capitalize()} intern application: {intern.name}"
)

    return redirect(url_for("admin_requests"))
@app.route("/admin/intern/<int:intern_id>/application")
@admin_required
def admin_view_application(intern_id):

    intern = User.query.filter_by(
        id=intern_id,
        role="intern"
    ).first_or_404()

    profile = InternProfile.query.filter_by(
        user_id=intern.id
    ).first()

    if not profile:
        flash("Application not found.", "danger")
        return redirect(url_for("admin_intern_records"))

    return render_template(
        "admin_view_application.html",
        intern=intern,
        profile=profile
    )

from flask import send_file
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
import io

@app.route("/admin/intern/<int:intern_id>/download-application")
def admin_download_application(intern_id):

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    intern = User.query.get_or_404(intern_id)
    profile = InternProfile.query.filter_by(user_id=intern.id).first()

    if not profile:
        flash("Application not found.", "danger")
        return redirect(url_for("admin_intern_records"))

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer)
    elements = []

    styles = getSampleStyleSheet()

    elements.append(Paragraph("<b>Intern Application Details</b>", styles["Title"]))
    elements.append(Spacer(1, 0.3 * inch))

    data_lines = [
        f"Name: {intern.name}",
        f"Email: {intern.email}",
        f"Course: {profile.course}",
        f"Year: {profile.year}",
        f"Semester: {profile.semester}",
        f"College: {profile.college}",
        f"Phone: {profile.phone}",
        f"Internship Duration: {profile.start_date} to {profile.end_date}",
        f"Objective: {profile.objective}",
        f"Application Status: {profile.application_status}"
    ]

    for line in data_lines:
        elements.append(Paragraph(line, styles["Normal"]))
        elements.append(Spacer(1, 0.2 * inch))

    doc.build(elements)

    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{intern.name}_application.pdf",
        mimetype="application/pdf"
    )
    
@app.route("/intern/upload-report", methods=["POST"])
def intern_upload_report():

    if "user_id" not in session:
        return redirect(url_for("intern_login"))

    intern = User.query.get(session["user_id"])
    profile = InternProfile.query.filter_by(user_id=intern.id).first()

    if not profile or not profile.report_upload_enabled:
        flash("Report upload not allowed.", "danger")
        return redirect(url_for("intern_dashboard"))

    file = request.files.get("report")

    if file and file.filename.endswith(".pdf"):

        filename = f"intern_{intern.id}_report.pdf"
        filepath = os.path.join("uploads", filename)

        file.save(filepath)

        # ✅ Save path in DB
        profile.report_file = filepath
        profile.report_upload_enabled = False  # disable after submission
        db.session.commit()

        flash("Report submitted successfully.", "success")

    return redirect(url_for("intern_dashboard"))

from flask import send_file
import os

@app.route("/admin/intern/<int:intern_id>/view-report")
def admin_view_report(intern_id):

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    intern = User.query.get_or_404(intern_id)
    profile = InternProfile.query.filter_by(user_id=intern.id).first()

    if not profile or not profile.report_file:
        flash("Report not found.", "danger")
        return redirect(url_for("admin_intern_days", intern_id=intern.id))

    # Build correct full path
    report_path = os.path.join(
        app.config["UPLOAD_FOLDER"],   # should be "static/reports"
        profile.report_file
    )

    # Make sure file actually exists
    if not os.path.exists(report_path):
        flash("Report file missing from server.", "danger")
        return redirect(url_for("admin_intern_days", intern_id=intern.id))

    return send_file(report_path)

@app.route("/admin/intern/<int:intern_id>/delete-report", methods=["POST"])
def admin_delete_report(intern_id):

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    intern = User.query.get_or_404(intern_id)
    profile = InternProfile.query.filter_by(user_id=intern.id).first()

    if not profile:
        flash("Intern profile not found.", "danger")
        return redirect(url_for("admin_intern_days", intern_id=intern.id))

    # ------------------------------
    # 1️⃣ Delete Report File (if exists)
    # ------------------------------
    if profile.report_file:
        report_path = os.path.join(app.config["UPLOAD_FOLDER"], profile.report_file)
        if os.path.exists(report_path):
            os.remove(report_path)

    profile.report_file = None

    # ------------------------------
    # 2️⃣ If certificate exists → delete it
    # ------------------------------
    if profile.certificate_file:
        cert_path = os.path.join(app.config["CERTIFICATE_FOLDER"], profile.certificate_file)
        if os.path.exists(cert_path):
            os.remove(cert_path)

    # Reset certificate related fields
    profile.certificate_file = None
    profile.survey_completed = False
    profile.survey_feedback = None
    profile.is_completed = False
    profile.certificate_requested = False
    profile.completion_date = None

    # Keep upload enabled so intern can re-upload
    profile.report_upload_enabled = True

    db.session.commit()

    flash("Report and associated certificate deleted successfully.", "warning")
    return redirect(url_for("admin_intern_days", intern_id=intern.id))
import math

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, HRFlowable
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from datetime import datetime
import os
import qrcode
import pytz
import random
import string


def generate_certificate(intern, profile):

    # ================= SAFE CERTIFICATE ID =================
    if not profile.certificate_id:

        current_year = datetime.now().year

    # Get highest serial issued this year
        last_serial = db.session.query(
            db.func.max(InternProfile.certificate_serial)
    ).filter(
            InternProfile.issued_at != None,
            db.extract('year', InternProfile.issued_at) == current_year
    ).scalar()

        if last_serial is None:
            new_serial = 1
        else:
            new_serial = last_serial + 1

        profile.certificate_serial = new_serial

        serial_formatted = f"{new_serial:03d}"

        random_suffix = ''.join(
            random.choices(string.ascii_uppercase + string.digits, k=6)
    )

        profile.certificate_id = (
            f"BBP-{current_year}-INT-{serial_formatted}-{random_suffix}"
    )

        ist = pytz.timezone('Asia/Kolkata')
        profile.issued_at = datetime.now(ist)

        db.session.commit()



    certificate_id = profile.certificate_id
    issued_at = profile.issued_at

    file_name = f"{certificate_id}.pdf"
    file_path = os.path.join(app.config["CERTIFICATE_FOLDER"], file_name)

    doc = SimpleDocTemplate(
        file_path,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50
    )

    elements = []
    styles = getSampleStyleSheet()

    # ================= STYLES =================

    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        alignment=1,
        fontName='CinzelBlack',
        fontSize=30,
        leading=36,
        textColor=colors.HexColor("#1B4332"),
        spaceAfter=20
    )

    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        alignment=1,
        fontName='PatrickRegular',
        fontSize=15,
        leading=22
    )

    name_style = ParagraphStyle(
        'NameStyle',
        parent=styles['Normal'],
        alignment=1,
        fontName='PatrickRegular',
        fontSize=26,
        leading=30,
        spaceAfter=8
    )

    bottom_left_style = ParagraphStyle(
        'BottomLeftStyle',
        parent=styles['Normal'],
        alignment=0,
        fontName='PatrickRegular',
        fontSize=11,
        textColor=colors.HexColor("#444444"),
        leading=14
    )

    bottom_right_style = ParagraphStyle(
        'BottomRightStyle',
        parent=styles['Normal'],
        alignment=1,
        fontName='PatrickRegular',
        fontSize=10,
        textColor=colors.HexColor("#777777"),
        leading=12
    )

    # ================= LOGO =================

    logo_path = os.path.join("static", "logo.png")
    if os.path.exists(logo_path):
        try:
            logo = Image(logo_path)
            logo.drawHeight = 1.0 * inch
            logo.drawWidth = logo.drawHeight * logo.imageWidth / logo.imageHeight
            logo.hAlign = "CENTER"
            elements.append(logo)
        except:
            pass

    elements.append(Spacer(1, 0.2 * inch))

    # ================= TITLE =================

    elements.append(
        Paragraph("CERTIFICATE OF APPRECIATION", title_style)
    )

    elements.append(
        HRFlowable(
            width=doc.width * 0.75,
            thickness=1.3,
            color=colors.HexColor("#C6A756"),
            hAlign='CENTER',
            spaceAfter=18
        )
    )

    # ================= BODY =================

    elements.append(Paragraph("This Certificate is proudly presented to", body_style))
    elements.append(Spacer(1, 0.15 * inch))

    elements.append(Paragraph(f"<b>{intern.name}</b>", name_style))
    elements.append(Spacer(1, 0.15 * inch))

    text = f"""
    for taking part in <b>Internship 2026</b>, hosted by Bannerghatta Biological Park, Bengaluru<br/>
    between {profile.start_date.strftime('%d %B %Y')} and {profile.end_date.strftime('%d %B %Y')}.<br/><br/>
    Congratulations on becoming <b>Junior Wildlife Ambassador</b> and we urge you to continue your efforts towards biodiversity conservation.
    """

    elements.append(Paragraph(text, body_style))
    elements.append(Spacer(1, 0.6 * inch))

    # ================= QR BLOCK =================
    # ===== GENERATE QR CODE =====

    verify_url = url_for("verify_certificate", certificate_id=certificate_id, _external=True)

    qr = qrcode.make(verify_url)

    qr_path = os.path.join(
    app.config["CERTIFICATE_FOLDER"],
    f"{certificate_id}_qr.png"
)

    qr.save(qr_path)

    qr_img = Image(qr_path)
    qr_img.drawHeight = 1.4 * inch
    qr_img.drawWidth = 1.4 * inch

    qr_block = Table([
        [qr_img],
        [Spacer(1, 0.05 * inch)],
        [Paragraph("Scan to verify", bottom_right_style)]
])

    qr_block.setStyle(TableStyle([
    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
]))

# ================= LEFT META BLOCK =================

    left_meta = Paragraph(
        f"Certificate ID: {certificate_id}<br/>"
        f"Issued On: {issued_at.strftime('%d %B %Y, %H:%M:%S IST')}",
        bottom_left_style
)

# ================= FINAL BOTTOM TABLE =================

    bottom_table = Table(
        [[left_meta, qr_block]],
        colWidths=[doc.width * 0.55, doc.width * 0.45]
)

    bottom_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),  # THIS FIXES ALIGNMENT
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
]))

    elements.append(bottom_table)


    # ================= BUILD =================

    doc.build(elements, onFirstPage=add_border)

    return file_name


def add_border(canvas, doc):
    width, height = A4

    # ================= TOGGLE CONTROLS =================
    USE_HORNBILL = True
    USE_LEAF_CORNERS = False
    USE_RADIAL_TEXTURE = False
    # Set any of these to False to disable

    # ===================================================
    # 1️⃣ RADIAL DIGITAL TEXTURE
    # ===================================================
    if USE_RADIAL_TEXTURE:
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#EDEDED"))
        canvas.setLineWidth(0.3)

        center_x = width / 2
        center_y = height * 0.65

        for i in range(0, 360, 10):
            canvas.line(
                center_x,
                center_y,
                center_x + 220 * math.cos(math.radians(i)),
                center_y + 220 * math.sin(math.radians(i))
            )

        canvas.restoreState()

    # ===================================================
    # 2️⃣ ELEPHANT WATERMARK (CENTER)
    # ===================================================
    if USE_HORNBILL:
        elephant_path = os.path.join("static", "hornbill_outline.png")

        if os.path.exists(elephant_path):
            canvas.saveState()
            canvas.setFillAlpha(0.04)  # very subtle
            watermark_width = width * 0.55

            canvas.drawImage(
                elephant_path,
                (width - watermark_width) / 2,
                (height - watermark_width) / 2,
                width=watermark_width,
                height=watermark_width,
                preserveAspectRatio=True,
                mask='auto'
            )

            canvas.restoreState()

    # ===================================================
    # 3️⃣ LEAF CORNER ACCENTS
    # ===================================================
    if USE_LEAF_CORNERS:
        leaf_tl = os.path.join("static", "leaf_top_left.png")
        leaf_br = os.path.join("static", "leaf_bottom_right.png")

        canvas.saveState()
        canvas.setFillAlpha(0.06)

        if os.path.exists(leaf_tl):
            canvas.drawImage(
                leaf_tl,
                35,
                height - 160,
                width=120,
                preserveAspectRatio=True,
                mask='auto'
            )

        if os.path.exists(leaf_br):
            canvas.drawImage(
                leaf_br,
                width - 155,
                35,
                width=120,
                preserveAspectRatio=True,
                mask='auto'
            )

        canvas.restoreState()

    # ===================================================
    # BORDERS (Always On)
    # ===================================================

    canvas.setStrokeColor(colors.HexColor("#1B4332"))
    canvas.setLineWidth(3)
    canvas.rect(30, 30, width - 60, height - 60)

    canvas.setStrokeColor(colors.HexColor("#C6A756"))
    canvas.setLineWidth(1)
    canvas.rect(40, 40, width - 80, height - 80)






from datetime import datetime

@app.route("/verify/<certificate_id>")
def verify_certificate(certificate_id):

    profile = InternProfile.query.filter_by(
        certificate_id=certificate_id
    ).first()

    if not profile:
        return render_template("verify_invalid.html"), 404

    return render_template(
        "verify_certificate.html",
        intern=profile.user,
        profile=profile,
        verified_at=datetime.utcnow()
    )


from datetime import date
import os

@app.route("/admin/intern/<int:intern_id>/complete", methods=["POST"])
@admin_required
def complete_internship(intern_id):

    intern = User.query.get_or_404(intern_id)
    profile = InternProfile.query.filter_by(user_id=intern.id).first()

    if not profile:
        flash("Intern profile not found.", "danger")
        return redirect(url_for("admin_intern_days", intern_id=intern.id))

    # 🚫 Prevent duplicate completion
    if profile.is_completed:
        flash("Internship has already been completed.", "warning")
        return redirect(url_for("admin_intern_days", intern_id=intern.id))

    # 🔍 Fetch all internship logs
    logs = InternDailyLog.query.filter_by(intern_id=intern.id).all()

    if not logs:
        flash("No internship days have been recorded.", "danger")
        return redirect(url_for("admin_intern_days", intern_id=intern.id))

    # ✅ Validate every log
    for log in logs:

        if not log.submitted:
            flash("All internship days must be submitted before completion.", "danger")
            return redirect(url_for("admin_intern_days", intern_id=intern.id))

        if not log.logout_time:
            flash("All internship days must be properly closed (logout required) before completion.", "danger")
            return redirect(url_for("admin_intern_days", intern_id=intern.id))

    # 📄 Validate final report
    if not profile.report_file:
        flash("Final internship report has not been uploaded.", "danger")
        return redirect(url_for("admin_intern_days", intern_id=intern.id))

    # 🎓 Generate certificate
    certificate_filename = generate_certificate(intern, profile)

    # ✅ Update profile status
    profile.internship_status = "completed"
    profile.is_completed = True
    profile.certificate_file = certificate_filename
    profile.report_upload_enabled = False
    profile.completion_date = datetime.now()

    db.session.commit()

    flash("Internship completed successfully. Certificate generated.", "success")

    return redirect(url_for("admin_intern_days", intern_id=intern.id))


from flask import send_file, abort
import os

@app.route("/intern/certificate")
def intern_certificate():

    if "user_id" not in session or session.get("role") != "intern":
        return redirect(url_for("intern_login"))

    profile = InternProfile.query.filter_by(
        user_id=session["user_id"]
    ).first()

    # No profile or not completed
    if not profile or not profile.is_completed or not profile.certificate_file:
        abort(403)

    certificate_path = os.path.join(
    app.root_path,
    "storage",
    "certificates",
    profile.certificate_file
)


    if not os.path.exists(certificate_path):
        abort(404)

    return send_file(certificate_path)
@app.route("/admin/intern/<int:intern_id>/certificate")
def admin_download_certificate(intern_id):

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    profile = InternProfile.query.filter_by(user_id=intern_id).first_or_404()

    certificate_folder = os.path.join(app.root_path, "storage", "certificates")
    certificate_path = os.path.join(certificate_folder, profile.certificate_file)

    if not os.path.exists(certificate_path):
        abort(404)

    return send_file(certificate_path, as_attachment=True)

@app.route("/intern/request-certificate", methods=["POST"])
def request_certificate():

    if "user_id" not in session or session.get("role") != "intern":
        return redirect(url_for("intern_login"))

    user = User.query.get(session["user_id"])
    profile = user.intern_profile

    if not profile.survey_completed:
        flash("Please complete the survey before requesting certificate.", "danger")
        return redirect(url_for("intern_dashboard"))

    if profile.certificate_requested:
        flash("Certificate already requested.", "warning")
        return redirect(url_for("intern_dashboard"))

    profile.certificate_requested = True
    db.session.commit()

    flash("Certificate request sent successfully.", "success")
    return redirect(url_for("intern_dashboard"))


@app.route("/admin/delete-certificate/<int:intern_id>", methods=["POST"])
def delete_certificate(intern_id):

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    profile = InternProfile.query.filter_by(user_id=intern_id).first_or_404()

    if profile.certificate_file:
        file_path = os.path.join(app.config["CERTIFICATE_FOLDER"], profile.certificate_file)

        if os.path.exists(file_path):
            os.remove(file_path)

    profile.is_completed = False
    profile.certificate_file = None
    profile.certificate_requested = False

    db.session.commit()

    flash("Certificate removed successfully.", "success")
    return redirect(url_for("admin_intern_days", intern_id=intern_id))

@app.route("/admin/interns/active")
@admin_required
def active_interns():

    interns = InternProfile.query.filter_by(internship_status="active").all()
    return render_template("admin_active_interns.html", interns=interns)

@app.route("/intern/post-survey", methods=["GET", "POST"])
def intern_post_survey():

    if "user_id" not in session or session.get("role") != "intern":
        return redirect(url_for("intern_login"))

    user = User.query.get(session["user_id"])
    profile = user.intern_profile

    if request.method == "POST":

        if profile.survey_completed:
            flash("Survey already submitted.", "warning")
            return redirect(url_for("intern_dashboard"))

        rating = request.form.get("rating")
        feedback = request.form.get("feedback")

        if not rating or not feedback:
            flash("Please complete all fields.", "danger")
            return redirect(url_for("intern_post_survey"))

        profile.survey_rating = int(rating)
        profile.survey_feedback = feedback
        profile.survey_completed = True

        db.session.commit()

        flash("Survey submitted successfully.", "success")
        return redirect(url_for("intern_dashboard"))

    return render_template("intern_post_survey.html", profile=profile)




@app.route("/admin/intern/<int:intern_id>/delete-survey", methods=["POST"])
def admin_delete_survey(intern_id):

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    intern = User.query.get_or_404(intern_id)
    profile = intern.intern_profile

    profile.survey_rating = None
    profile.survey_feedback = None
    profile.survey_completed = False

    db.session.commit()

    flash("Survey deleted successfully.", "success")
    return redirect(url_for("admin_intern_days", intern_id=intern.id))







@app.route("/admin/close-internship/<int:user_id>", methods=["POST"])
@admin_required
def close_internship(user_id):

    user = User.query.get_or_404(user_id)
    profile = user.intern_profile

    if not profile:
        flash("Intern profile not found.", "danger")
        return redirect(request.referrer)

    # ✅ CHECK: Are all daily logs properly closed?
    incomplete_logs = InternDailyLog.query.filter(
        InternDailyLog.intern_id == user_id,
        InternDailyLog.logout_time == None
    ).all()

    if incomplete_logs:
        flash(
            "Cannot close internship. Some daily records do not have logout time.",
            "danger"
        )
        return redirect(request.referrer)

    # ✅ If all good → close internship
    profile.is_completed = True
    profile.report_upload_enabled = False
    db.session.commit()

    flash("Internship closed successfully.", "success")
    return redirect(request.referrer)


@app.route("/admin/intern/<int:intern_id>/reopen", methods=["POST"])
def reopen_internship(intern_id):

    if not session.get("admin_logged_in"):
        return redirect(url_for("admin_login"))

    intern = User.query.get_or_404(intern_id)
    profile = intern.intern_profile

    if profile:
        profile.is_completed = False
        profile.completion_date = None

        db.session.commit()

        flash("Internship reopened successfully.", "success")

    return redirect(url_for("admin_intern_days", intern_id=intern.id))

@app.route("/enroll-programs")
def enroll_programs():
    return render_template("enroll_programs.html")


@app.route("/zoo-club")
def zoo_club_programs():
    return render_template("zoo_club.html")




@app.route("/summer-camp")
def summer_camp_programs():
    return render_template("summer_camp.html")
from flask import request


from flask import request, session

def log_action(section, action, target_type=None, target_id=None, description=None):

    log = AuditLog(
        actor_id=session.get("admin_id"),
        actor_role=session.get("admin_role", "admin"),
        section=section,
        action=action,
        target_type=target_type,
        target_id=target_id,
        description=description,
        ip_address=request.remote_addr
    )

    db.session.add(log)
    db.session.commit()

# =========================
# RUN
# =========================
if __name__ == "__main__":
    app.run(debug=True)
